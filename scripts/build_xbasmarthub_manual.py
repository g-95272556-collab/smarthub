from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "deliverables"
DOCX_PATH = OUTPUT_DIR / "manual-penggunaan-xbasmarthub-sk-kiandongo.docx"
LOGO_PATH = ROOT / "assets" / "sk-kiandongo-logo.png"

NAVY = RGBColor(20, 36, 68)
BLUE = RGBColor(48, 98, 168)
GOLD = RGBColor(214, 173, 38)
TEXT = RGBColor(33, 37, 41)
MUTED = RGBColor(98, 108, 127)
LIGHT = RGBColor(244, 247, 251)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color="D8E1EC", size=8):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
      borders = OxmlElement("w:tblBorders")
      tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
      el = borders.find(qn(f"w:{edge}"))
      if el is None:
        el = OxmlElement(f"w:{edge}")
        borders.append(el)
      el.set(qn("w:val"), "single")
      el.set(qn("w:sz"), str(size))
      el.set(qn("w:space"), "0")
      el.set(qn("w:color"), color)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in (
        ("Title", 24, NAVY),
        ("Heading 1", 16, NAVY),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11.5, NAVY),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Manual Penggunaan xbasmarthub | SK Kiandongo | Halaman ")
    run.font.name = "Aptos"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED
    add_page_number(p)


def add_cover(doc):
    if LOGO_PATH.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(str(LOGO_PATH), width=Inches(1.35))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("MANUAL PENGGUNAAN\nWEBAPP XBasSmartHub")
    r.font.name = "Aptos Display"
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = subtitle.add_run("Panduan mesra guru untuk penggunaan harian di sekolah")
    rs.font.name = "Aptos"
    rs.font.size = Pt(12)
    rs.font.color.rgb = MUTED

    info = doc.add_table(rows=4, cols=2)
    info.autofit = False
    info.columns[0].width = Inches(2.0)
    info.columns[1].width = Inches(4.8)
    rows = [
        ("Sekolah", "SK Kiandongo"),
        ("Aplikasi", "SmartSchoolHub v2.0 / xbasmarthub"),
        ("Sasaran", "Guru, guru bertugas, penyelaras program dan pentadbir sekolah"),
        ("Tujuan", "Menyelaras kehadiran, OPR, laporan harian, kokurikulum dan operasi sekolah dalam satu webapp"),
    ]
    for idx, (label, value) in enumerate(rows):
        info.cell(idx, 0).text = label
        info.cell(idx, 1).text = value
        set_cell_shading(info.cell(idx, 0), "E8EEF8")
    set_table_borders(info, color="C9D6E7", size=10)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = note.add_run(
        "Cadangan penggunaan semasa taklimat: buka webapp pada projektor dan ikut langkah dalam manual ini mengikut modul."
    )
    rn.italic = True
    rn.font.color.rgb = BLUE
    doc.add_page_break()


def add_highlight_box(doc, title, bullets):
    table = doc.add_table(rows=len(bullets) + 1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.45)
    head = table.cell(0, 0)
    head.text = title
    set_cell_shading(head, "1F3A5F")
    p = head.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if p.runs:
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.bold = True
    for i, bullet in enumerate(bullets, start=1):
        cell = table.cell(i, 0)
        set_cell_shading(cell, "F7FAFD")
        para = cell.paragraphs[0]
        para.style = doc.styles["Normal"]
        para.paragraph_format.left_indent = Inches(0.15)
        run = para.add_run("• " + bullet)
        run.font.color.rgb = TEXT
    set_table_borders(table, color="D7E2EF", size=8)
    doc.add_paragraph()


def add_section_title(doc, title, subtitle=None):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(title)
    if subtitle:
        ps = doc.add_paragraph()
        rs = ps.add_run(subtitle)
        rs.font.size = Pt(10)
        rs.font.color.rgb = MUTED


def add_numbered_steps(doc, items):
    for idx, item in enumerate(items, start=1):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.05)
        para.paragraph_format.space_after = Pt(4)
        num = para.add_run(f"{idx}. ")
        num.bold = True
        num.font.color.rgb = NAVY
        body = para.add_run(item)
        body.font.color.rgb = TEXT


def add_two_col_table(doc, title, rows):
    doc.add_paragraph(title, style="Heading 2")
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.95)
    table.columns[1].width = Inches(4.5)
    headers = ["Perkara", "Penerangan"]
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = text
        set_cell_shading(cell, "E8EEF8")
    for row_idx, (left, right) in enumerate(rows, start=1):
        table.cell(row_idx, 0).text = left
        table.cell(row_idx, 1).text = right
    set_table_borders(table, color="D0DBEA", size=8)
    doc.add_paragraph()


def build_manual():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_cover(doc)

    add_section_title(doc, "1. Pengenalan Ringkas", "Webapp ini menghimpunkan urusan harian guru dalam satu tempat yang lebih cepat dan tersusun.")
    doc.add_paragraph(
        "XBasSmartHub dibina sebagai webapp sekolah untuk membantu guru merekod kehadiran, mengurus laporan, merekod OPR, "
        "menghantar pelaporan kokurikulum dan menyemak maklumat sekolah tanpa perlu berpindah-pindah platform."
    )
    add_highlight_box(doc, "Apa yang guru perlu tahu dahulu", [
        "Aplikasi ini dibuka melalui pelayar web pada telefon, tablet atau komputer.",
        "Sebahagian fungsi memerlukan log masuk akaun Google yang dibenarkan oleh pentadbir.",
        "Modul guru yang paling kerap digunakan ialah Kehadiran Guru, OPR, Laporan Guru Bertugas, Takwim dan Pelaporan Kokum.",
        "Modul Konfigurasi dan Notifikasi tertentu adalah khas untuk pentadbir sekolah.",
    ])

    add_section_title(doc, "2. Persediaan Sebelum Menggunakan Webapp")
    add_numbered_steps(doc, [
        "Pastikan sambungan internet stabil dan pelayar web dikemas kini.",
        "Buka pautan rasmi webapp xbasmarthub yang diberikan oleh pihak sekolah.",
        "Log masuk menggunakan akaun Google yang telah didaftarkan dalam sistem.",
        "Benarkan akses lokasi apabila diminta jika mahu menggunakan daftar hadir berasaskan GPS.",
        "Jika butang Pasang Aplikasi dipaparkan, guru digalakkan memasang PWA supaya capaian lebih pantas seperti aplikasi sebenar.",
    ])
    add_two_col_table(doc, "Keperluan minimum penggunaan", [
        ("Peranti", "Telefon Android/iPhone, tablet, komputer riba atau desktop"),
        ("Pelayar", "Google Chrome, Microsoft Edge atau pelayar moden yang menyokong PWA"),
        ("Log masuk", "Akaun Google guru yang telah didaftarkan"),
        ("Lokasi", "Perlu diaktifkan untuk kehadiran GPS"),
    ])

    add_section_title(doc, "3. Kenali Paparan Utama dan Menu")
    doc.add_paragraph(
        "Selepas log masuk, guru akan melihat panel sisi kiri yang memaparkan modul utama. Gunakan menu ini untuk berpindah modul "
        "tanpa menutup halaman."
    )
    add_two_col_table(doc, "Modul utama untuk guru", [
        ("Papan Pemuka", "Paparan ringkas status kehadiran, takwim, notifikasi dan ringkasan modul."),
        ("Kehadiran Guru", "Daftar masuk, punch out, semakan status lokasi dan rekod kehadiran harian."),
        ("OPR", "Merekod laporan ringkas aktiviti/program dan mencetak atau eksport laporan."),
        ("Laporan Guru Bertugas", "Catatan mingguan guru bertugas termasuk disiplin, kebersihan dan rumusan."),
        ("Pelaporan Kokum", "Rekod aktiviti kokurikulum, unit, kelab, sukan, kehadiran dan status pelaporan."),
        ("Takwim", "Semakan aktiviti sekolah dan perancangan program."),
        ("Lembaran Kerja", "Ruang sokongan pengurusan bahan atau kerja guru mengikut penggunaan sekolah."),
    ])

    add_section_title(doc, "4. Cara Menggunakan Modul Kehadiran Guru")
    add_highlight_box(doc, "Aliran paling cepat setiap pagi", [
        "Buka modul Kehadiran Guru.",
        "Semak status lokasi atau tekan Semak GPS jika perlu.",
        "Tekan Daftar Hadir Sekarang semasa dalam tempoh hadir.",
        "Semak status berjaya direkod sebelum keluar dari modul.",
        "Gunakan Punch-Out Sekarang atau rekod keluar mengikut aliran sekolah pada akhir sesi.",
    ])
    add_two_col_table(doc, "Fungsi dalam modul Kehadiran Guru", [
        ("Segar Sekarang", "Muat semula paparan kehadiran guru."),
        ("Rekod Manual Admin", "Digunakan oleh pentadbir untuk merekodkan kehadiran staf secara manual."),
        ("Daftar Keluar Manual", "Merekod punch out secara manual jika diperlukan."),
        ("Daftar Manual", "Pilihan rekod manual apabila akses GPS tidak sesuai atau tidak tersedia."),
    ])
    doc.add_paragraph(
        "Jika sistem memaparkan mesej seperti tetingkap auto daftar hadir telah tamat, guru masih boleh menyemak keadaan semasa dan "
        "mengikut arahan pentadbir untuk rekod manual sekiranya perlu."
    )

    add_section_title(doc, "5. Cara Menggunakan Modul OPR")
    add_numbered_steps(doc, [
        "Buka modul OPR selepas selesai program atau aktiviti.",
        "Isi maklumat asas seperti nama program, tarikh, tempat, peserta dan penyedia laporan.",
        "Lengkapkan objektif, aktiviti dijalankan, kekuatan/kejayaan serta kelemahan/cadangan.",
        "Muat naik sekurang-kurangnya gambar aktiviti yang diperlukan mengikut ketetapan sekolah.",
        "Semak semula semua maklumat sebelum menekan butang Cetak / Eksport PDF.",
    ])
    add_highlight_box(doc, "Amalan baik semasa menyediakan OPR", [
        "Gunakan bahasa rasmi dan ringkas.",
        "Pastikan gambar jelas, relevan dan mewakili aktiviti sebenar.",
        "Semak ejaan nama program, tarikh dan bilangan peserta sebelum cetak.",
    ])

    add_section_title(doc, "6. Cara Menggunakan Laporan Guru Bertugas")
    doc.add_paragraph(
        "Modul ini digunakan oleh guru bertugas untuk merekod pemerhatian harian atau mingguan seperti kebersihan, disiplin, "
        "kehadiran, cuaca dan rumusan tindakan."
    )
    add_numbered_steps(doc, [
        "Pilih minggu atau tarikh bertugas yang berkaitan.",
        "Isi semua seksyen penting secara lengkap dan ringkas.",
        "Gunakan rumusan yang formal supaya mudah dibentang kepada pentadbir.",
        "Semak sekali lagi sebelum simpan atau cetak laporan.",
    ])

    add_section_title(doc, "7. Cara Menggunakan Pelaporan Kokum")
    doc.add_paragraph(
        "Modul Pelaporan Kokum membantu guru penasihat dan penyelaras merekod aktiviti kokurikulum dengan lebih kemas."
    )
    add_two_col_table(doc, "Maklumat yang biasanya direkod", [
        ("Unit/Program", "Badan beruniform, kelab/persatuan atau sukan/permainan"),
        ("Kehadiran", "Bilangan atau senarai murid yang terlibat"),
        ("Aktiviti", "Ringkasan aktiviti, latihan, pertandingan atau pelaksanaan"),
        ("Status", "Pencapaian, tindakan susulan atau notifikasi kepada penyelaras/PK Kokum"),
    ])

    add_section_title(doc, "8. Takwim, Notifikasi dan Modul Admin")
    doc.add_paragraph(
        "Guru boleh menggunakan modul Takwim untuk menyemak aktiviti sekolah. Modul Notifikasi dan kebanyakan tetapan Konfigurasi "
        "dikawal oleh pentadbir. Semasa taklimat, jelaskan kepada guru bahawa mereka hanya perlu fokus pada modul yang menyokong tugas harian."
    )
    add_highlight_box(doc, "Perkara penting tentang akses", [
        "Guru biasa: fokus pada Kehadiran, OPR, Laporan, Takwim dan Pelaporan Kokum.",
        "Pentadbir: mengurus konfigurasi, notifikasi, data operasi dan bantuan rekod manual.",
        "Jika menu tertentu tidak dipaparkan, itu biasanya kerana akses mengikut peranan pengguna.",
    ])

    add_section_title(doc, "9. Penyelesaian Masalah Ringkas")
    add_two_col_table(doc, "Masalah lazim dan tindakan pantas", [
        ("Tidak boleh log masuk", "Semak sama ada akaun Google yang digunakan telah didaftarkan dan cuba log masuk semula."),
        ("Butang kehadiran tidak berjaya", "Semak internet, benarkan lokasi, tekan Segar Sekarang atau maklum kepada pentadbir untuk rekod manual."),
        ("GPS tidak tepat", "Aktifkan lokasi ketepatan tinggi dan cuba Semak GPS semula."),
        ("Data tidak dikemas kini", "Muat semula modul atau log keluar dan log masuk semula."),
        ("Tidak nampak modul tertentu", "Akses mungkin khusus kepada pentadbir atau peranan tertentu."),
    ])

    add_section_title(doc, "10. Cadangan Aliran Pembentangan Kepada Guru")
    add_numbered_steps(doc, [
        "Mulakan dengan tujuan webapp: satu tempat untuk urusan harian guru.",
        "Tunjukkan cara log masuk dan pemasangan aplikasi pada telefon.",
        "Demonstrasi aliran pagi: semak GPS, daftar hadir dan semak status berjaya.",
        "Buka modul OPR dan tunjuk cara isi laporan selepas program.",
        "Tunjukkan modul guru bertugas, kokum dan takwim secara ringkas.",
        "Akhiri dengan tips masalah lazim dan siapa yang perlu dihubungi jika berlaku isu.",
    ])

    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = closing.add_run("Manual ini disediakan sebagai panduan penggunaan asas semasa taklimat dan penggunaan harian guru di sekolah.")
    rc.italic = True
    rc.font.color.rgb = MUTED

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_manual()
    print(path)
